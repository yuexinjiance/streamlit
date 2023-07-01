import os
from tkinter import messagebox
# 导入在线数据
from get_new_data import get_new_data
# 安装修改word文档的库
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

# 获取在线数据
product_json, factory_list = get_new_data()


def deal_input_number(user_input):
    return_list = []
    user_input_list = user_input.strip().split(' ')
    for item in user_input_list:
        if '-' in item:
            first_num = int(float(item.split('-')[0]))
            second_num = int(float(item.split('-')[1]))
            for i in range(first_num, second_num + 1):
                return_list.append("{:03d}".format(i))
        else:
            return_list.append("{:03d}".format(int(float(item))))
    return return_list


def get_all_problem_files_name(company_name, user_input_num):
    problem_nums_list = deal_input_number(user_input_num)
    all_files = os.listdir(f"{company_name}")
    all_problem_files_name = []
    for problem_num in problem_nums_list:
        for file in all_files:
            if problem_num == file.split('-')[1][-3:]:
                all_problem_files_name.append(f"{company_name}/{file}")
    return all_problem_files_name


def change_one_brand(problem_file_name, alert_factory, alert_type):
    factory_name = [item['fullname'] for item in product_json if item['name'] == alert_factory][0]
    dongzuozhi = [item['alarm'] for item in product_json if item['name'] == alert_factory][0]

    document = Document(problem_file_name)
    document.paragraphs[16].runs[7].text = alert_type
    document.paragraphs[22].runs[9].text = factory_name
    document.paragraphs[62].runs[3].text = str(dongzuozhi)

    document.save(problem_file_name)


# 更改品牌和型号
def change_all_brand(company_name, product_company, product_model, product_num_list):
    problem_files_name = get_all_problem_files_name(company_name, product_num_list)
    for problem_file_name in problem_files_name:
        change_one_brand(problem_file_name, product_company, product_model)
    messagebox.showinfo(title="修改完成", message="品牌已经修改完成，请到指定的公司文件夹查看")


# 更改单个探头为故障
def change_problem_file(problem_filename):
    document = Document(problem_filename)

    document.paragraphs[60].runs[2].text = "异常"  # 更改报警功能的值
    document.paragraphs[62].runs[3].text = "/"  # 更改报警动作值
    document.paragraphs[62].runs[4].text = ""
    document.paragraphs[65].runs[3].text = "/"  # 更改重复性
    document.paragraphs[65].runs[4].text = ""
    document.paragraphs[66].runs[3].text = "  /   "  # 更改响应时间
    document.paragraphs[66].runs[4].text = ""

    tables = document.tables
    table = tables[1]
    table.cell(1, 1).text = "/"
    table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, 1).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(1, 2).text = "/"
    table.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, 2).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(1, 3).text = "/"
    table.cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, 3).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(2, 1).text = "/"
    table.cell(2, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(2, 2).text = "/"
    table.cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 2).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(2, 3).text = "/"
    table.cell(2, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(2, 3).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(3, 1).text = "/"
    table.cell(3, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(3, 1).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(3, 2).text = "/"
    table.cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(3, 2).paragraphs[0].runs[0].font.size = Pt(12)
    table.cell(3, 3).text = "/"
    table.cell(3, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(3, 3).paragraphs[0].runs[0].font.size = Pt(12)

    document.save(problem_filename)


# 更改所有的故障探头
def change_all_problem_file(company_name, product_num_list):
    problem_files_name = get_all_problem_files_name(company_name, product_num_list)
    for problem_file_name in problem_files_name:
        change_problem_file(problem_file_name)
    messagebox.showinfo(title="修改完成", message="所有的异常探头报告已经修改完成，请到指定的公司文件夹查看")


# 查找指定文本内容在模板word文档中的对应位置
def find_text_in_word(text):
    filename = "test.docx"
    document = Document(filename)
    target_text = text
    # 遍历文档中的每个段落
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        # 遍历段落中的每个运行
        for run_index, run in enumerate(paragraph.runs):
            # 判断运行的文本是否与目标文本匹配
            if target_text in run.text:
                # 返回匹配到的段落序号和运行序号
                result = f"段落序号：{paragraph_index}，运行序号：{run_index}"
                return paragraph_index, run_index

        else:
            continue
    else:
        # 如果没有找到目标文本，将 result 设置为 None
        result = None


if __name__ == "__main__":
    change_all_brand("绍兴鑫沃工程有限公司", "珠海兴华", "XH-G800C", "7-10")
    # change_one_brand("test.docx", "珠海兴华", "XH-G800C")
    # change_all_problem_file("绍兴鑫沃工程有限公司", "2-5")
    # print(find_text_in_word("AEC2332"))
    # print(find_text_in_word("成都安可信电子股份有限公司"))
    # print(find_text_in_word("ankexindongzuo"))

