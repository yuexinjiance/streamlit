from .get_all_problem_files_name import get_all_problem_files_name
from docx import Document


# 更改品牌和型号
def change_all_brand(company_name, product_company_full_name, product_model, dongzuozhi, product_num_list):
    problem_files_name = get_all_problem_files_name(company_name, product_num_list)
    for problem_file_name in problem_files_name:
        document = Document(problem_file_name)
        document.paragraphs[16].runs[7].text = product_model
        document.paragraphs[22].runs[9].text = product_company_full_name
        document.paragraphs[62].runs[3].text = str(dongzuozhi)
        document.save(problem_file_name)
